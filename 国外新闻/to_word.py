# -*- coding: utf-8 -*-
"""
@Time ： 2023/2/8 3:36 下午
@Auth ： HaHa-Wa
@File ：to_word.py
@IDE ：PyCharm
"""
# 数据写入
from docx import Document
import pandas as pd
from docx.shared import Inches
# 创建对象
def save_word():
    document = Document()
    # 添加标题，其中'0'代表标题类型，共四种类型，具体可在Word的'开始'-'样式'中查看
    document.add_heading('Python 爬虫', 0)
    document.add_paragraph('标题')
    document.add_paragraph('时间')

    # 添加正文，设置'样式'-'明显引用'
    document.add_paragraph('样式-明显引用将卡上看见啊还是多看几哈上课打瞌睡的家')
    document.add_paragraph('\n\n\n')
    document.add_paragraph('标题')
    document.add_paragraph('时间')

    # 添加正文，设置'样式'-'明显引用'
    document.add_paragraph('样式-明显引用将卡上看见啊还是多看几哈上课打瞌睡的家')
    document.add_paragraph('\n\n\n')
    # 保存文件
    document.add_page_break()
    document.save('test.docx')

def read_pd(file_name):
    ret = pd.read_excel('file/{0}.xlsx'.format(file_name))
    document = Document()

    for x in ret.itertuples():
        print(x[2], x[3], x[4])
        # 添加标题，其中'0'代表标题类型，共四种类型，具体可在Word的'开始'-'样式'中查看
        document.add_heading(str(x[1])+x[2], 0)
        document.add_paragraph(x[3].replace('\n\n', ''))
        #
        # # 添加正文，设置'样式'-'明显引用'
        document.add_paragraph(x[4].replace('\n\n', ''))
        # document.add_paragraph('\n\n\n')

        # 保存文件
    document.add_page_break()
    document.save('{0}.docx'.format(file_name))

if __name__ == '__main__':
    file_names = ['ahram', 'akhbarelyom', 'almasryalyoum', 'youm7']
    for file_name in file_names:
        read_pd(file_name)
        # break
